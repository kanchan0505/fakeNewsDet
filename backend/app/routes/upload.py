import io
import os
from fastapi import APIRouter, UploadFile, File, HTTPException, status
import pypdf
import docx

router = APIRouter()

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}

@router.post("/extract-text")
async def extract_text(file: UploadFile = File(...)):
    # Get file extension
    filename = file.filename or ""
    _, ext = os.path.splitext(filename.lower())
    
    if ext not in SUPPORTED_EXTENSIONS:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Unsupported file type '{ext}'. Supported formats: PDF, DOCX, TXT, MD."
        )

    content = await file.read()
    
    if len(content) == 0:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="The uploaded file is empty."
        )
        
    extracted_text = ""
    
    try:
        if ext == ".pdf":
            try:
                pdf_reader = pypdf.PdfReader(io.BytesIO(content))
                pages_text = []
                for page in pdf_reader.pages:
                    text = page.extract_text()
                    if text:
                        pages_text.append(text)
                extracted_text = "\n".join(pages_text)
            except Exception as e:
                raise ValueError(f"Failed to parse PDF file: {str(e)}")
                
        elif ext == ".docx":
            try:
                doc = docx.Document(io.BytesIO(content))
                paragraphs = [p.text for p in doc.paragraphs]
                # Also extract text from tables if any
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            paragraphs.append(cell.text)
                extracted_text = "\n".join(paragraphs)
            except Exception as e:
                raise ValueError(f"Failed to parse Word document: {str(e)}")
                
        elif ext in {".txt", ".md"}:
            try:
                extracted_text = content.decode("utf-8")
            except UnicodeDecodeError:
                # Try fallback decoding
                try:
                    extracted_text = content.decode("latin-1")
                except Exception as e:
                    raise ValueError(f"Failed to decode text file: {str(e)}")
                    
        # Clean extracted text and check if it is empty after parsing
        extracted_text = extracted_text.strip()
        if not extracted_text:
            raise ValueError("No readable text could be extracted from the file.")
            
    except ValueError as val_err:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=str(val_err)
        )
    except Exception as err:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"An unexpected error occurred during text extraction: {str(err)}"
        )
        
    return {
        "text": extracted_text,
        "filename": filename
    }
